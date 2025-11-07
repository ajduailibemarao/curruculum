from __future__ import annotations

import io
from dataclasses import dataclass
from typing import Callable, Dict

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
from reportlab.platypus.tables import Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet

from ..models import ResumeData, TemplateMetadata


@dataclass
class TemplateDefinition:
    metadata: TemplateMetadata
    docx_renderer: Callable[[ResumeData], bytes]
    pdf_renderer: Callable[[ResumeData], bytes]


def _register_fonts() -> None:
    try:
        pdfmetrics.getFont("Helvetica")
    except KeyError:
        pdfmetrics.registerFont(TTFont("Helvetica", "Helvetica"))


def _create_header(document: Document, resume: ResumeData, color: RGBColor | None = None) -> None:
    name_paragraph = document.add_heading(resume.contact.full_name, level=1)
    name_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if color:
        for run in name_paragraph.runs:
            run.font.color.rgb = color

    info_paragraph = document.add_paragraph()
    info_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_parts = [
        part
        for part in [
            resume.contact.email,
            resume.contact.phone,
            resume.contact.location,
            resume.contact.linkedin,
            resume.contact.website,
        ]
        if part
    ]
    info_paragraph.add_run(" | ".join(contact_parts))


def _add_section_heading(document: Document, title: str, color: RGBColor | None = None) -> None:
    heading = document.add_heading(title, level=2)
    if color:
        for run in heading.runs:
            run.font.color.rgb = color


def _render_modern_blue_docx(resume: ResumeData) -> bytes:
    document = Document()
    _create_header(document, resume, RGBColor(0x1F, 0x4E, 0x79))

    if resume.professional_summary:
        _add_section_heading(document, "Resumo Profissional", RGBColor(0x1F, 0x4E, 0x79))
        document.add_paragraph(resume.professional_summary)

    if resume.experiences:
        _add_section_heading(document, "Experiência", RGBColor(0x1F, 0x4E, 0x79))
        for experience in resume.experiences:
            para = document.add_paragraph()
            para.add_run(f"{experience.role}").bold = True
            if experience.company:
                para.add_run(f" | {experience.company}")
            if experience.start_date or experience.end_date:
                para.add_run(f" ({experience.start_date or ''} - {experience.end_date or 'Atual'})")
            for highlight in experience.highlights:
                document.add_paragraph(highlight, style="List Bullet")

    if resume.educations:
        _add_section_heading(document, "Formação", RGBColor(0x1F, 0x4E, 0x79))
        for education in resume.educations:
            text = education.degree
            if education.institution:
                text += f" - {education.institution}"
            if education.summary:
                text += f" ({education.summary})"
            document.add_paragraph(text)

    if resume.skills:
        _add_section_heading(document, "Habilidades", RGBColor(0x1F, 0x4E, 0x79))
        document.add_paragraph(", ".join(resume.skills))

    if resume.projects:
        _add_section_heading(document, "Projetos", RGBColor(0x1F, 0x4E, 0x79))
        for project in resume.projects:
            text = project.name
            if project.link:
                text += f" - {project.link}"
            if project.description:
                text += f": {project.description}"
            document.add_paragraph(text)

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _render_classic_serif_docx(resume: ResumeData) -> bytes:
    document = Document()
    style = document.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)

    _create_header(document, resume)

    if resume.professional_summary:
        _add_section_heading(document, "Perfil", None)
        document.add_paragraph(resume.professional_summary)

    if resume.experiences:
        _add_section_heading(document, "Histórico Profissional", None)
        for experience in resume.experiences:
            para = document.add_paragraph()
            run = para.add_run(f"{experience.role} - {experience.company or ''}")
            run.bold = True
            if experience.start_date or experience.end_date:
                para.add_run(f" ({experience.start_date or ''} - {experience.end_date or 'Atual'})")
            for highlight in experience.highlights:
                document.add_paragraph(highlight, style="List Number")

    if resume.educations:
        _add_section_heading(document, "Educação", None)
        for education in resume.educations:
            para = document.add_paragraph()
            run = para.add_run(education.degree)
            run.bold = True
            if education.institution:
                para.add_run(f" - {education.institution}")
            if education.summary:
                document.add_paragraph(education.summary)

    if resume.skills:
        _add_section_heading(document, "Competências", None)
        document.add_paragraph(" • ".join(resume.skills))

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _render_minimalist_docx(resume: ResumeData) -> bytes:
    document = Document()
    section = document.sections[0]
    section.left_margin = section.right_margin = Pt(36)

    _create_header(document, resume, RGBColor(0x33, 0x33, 0x33))

    if resume.professional_summary:
        _add_section_heading(document, "Sobre", RGBColor(0x33, 0x33, 0x33))
        document.add_paragraph(resume.professional_summary)

    def add_two_column_list(title: str, items: Dict[str, str]) -> None:
        _add_section_heading(document, title, RGBColor(0x33, 0x33, 0x33))
        table = document.add_table(rows=1, cols=2)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Item"
        hdr_cells[1].text = "Detalhes"
        for key, value in items.items():
            row_cells = table.add_row().cells
            row_cells[0].text = key
            row_cells[1].text = value

    if resume.experiences:
        experience_items = {}
        for experience in resume.experiences:
            key = experience.role
            if experience.company:
                key += f" @ {experience.company}"
            timeframe = f"{experience.start_date or ''} - {experience.end_date or 'Atual'}"
            details = "; ".join(experience.highlights) or timeframe
            experience_items[key] = f"{timeframe} | {details}"
        add_two_column_list("Experiência", experience_items)

    if resume.educations:
        education_items = {}
        for education in resume.educations:
            key = education.degree
            details = education.institution or ""
            if education.summary:
                details += f" | {education.summary}"
            education_items[key] = details
        add_two_column_list("Formação", education_items)

    if resume.skills:
        _add_section_heading(document, "Competências", RGBColor(0x33, 0x33, 0x33))
        document.add_paragraph(", ".join(resume.skills))

    if resume.projects:
        _add_section_heading(document, "Projetos", RGBColor(0x33, 0x33, 0x33))
        for project in resume.projects:
            texto = project.name
            if project.link:
                texto += f" - {project.link}"
            if project.description:
                texto += f": {project.description}"
            document.add_paragraph(texto)

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _render_executive_gold_docx(resume: ResumeData) -> bytes:
    document = Document()
    style = document.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(12)

    section = document.sections[0]
    section.top_margin = Pt(36)
    section.bottom_margin = Pt(36)

    _create_header(document, resume, RGBColor(0xA5, 0x7C, 0x00))

    if resume.professional_summary:
        _add_section_heading(document, "Resumo Executivo", RGBColor(0xA5, 0x7C, 0x00))
        document.add_paragraph(resume.professional_summary)

    if resume.experiences:
        _add_section_heading(document, "Trajetória Profissional", RGBColor(0xA5, 0x7C, 0x00))
        for experience in resume.experiences:
            para = document.add_paragraph()
            role_run = para.add_run(experience.role)
            role_run.bold = True
            if experience.company:
                para.add_run(f" • {experience.company}")
            if experience.start_date or experience.end_date:
                para.add_run(f" ({experience.start_date or ''} - {experience.end_date or 'Atual'})")
            for highlight in experience.highlights:
                document.add_paragraph(highlight, style="List Bullet")

    if resume.skills:
        _add_section_heading(document, "Áreas de Expertise", RGBColor(0xA5, 0x7C, 0x00))
        document.add_paragraph(" | ".join(resume.skills))

    if resume.educations:
        _add_section_heading(document, "Formação Acadêmica", RGBColor(0xA5, 0x7C, 0x00))
        for education in resume.educations:
            text = education.degree
            if education.institution:
                text += f" - {education.institution}"
            if education.summary:
                text += f" ({education.summary})"
            document.add_paragraph(text)

    if resume.projects:
        _add_section_heading(document, "Resultados Relevantes", RGBColor(0xA5, 0x7C, 0x00))
        for project in resume.projects:
            texto = project.name
            if project.description:
                texto += f" — {project.description}"
            if project.link:
                texto += f" ({project.link})"
            document.add_paragraph(texto)

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _render_pdf_with_style(resume: ResumeData, theme_color: colors.Color, title: str) -> bytes:
    _register_fonts()
    buffer = io.BytesIO()

    doc = SimpleDocTemplate(buffer, pagesize=LETTER, topMargin=0.75 * inch, bottomMargin=0.75 * inch)
    doc.title = title
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name="HeadingColor", parent=styles["Heading2"], textColor=theme_color))

    story = []

    header_style = ParagraphStyle(
        name="Header",
        parent=styles["Title"],
        alignment=1,
        textColor=theme_color,
    )

    story.append(Paragraph(resume.contact.full_name, header_style))
    contact_parts = [
        part
        for part in [
            resume.contact.email,
            resume.contact.phone,
            resume.contact.location,
            resume.contact.linkedin,
            resume.contact.website,
        ]
        if part
    ]
    story.append(Paragraph(" | ".join(contact_parts), styles["Normal"]))
    story.append(Spacer(1, 12))

    if resume.professional_summary:
        story.append(Paragraph("Resumo Profissional", styles["Heading2"]))
        story.append(Paragraph(resume.professional_summary, styles["BodyText"]))
        story.append(Spacer(1, 12))

    if resume.experiences:
        story.append(Paragraph("Experiência", styles["Heading2"]))
        for experience in resume.experiences:
            header = f"<b>{experience.role}</b>"
            if experience.company:
                header += f" - {experience.company}"
            if experience.start_date or experience.end_date:
                header += f" ({experience.start_date or ''} - {experience.end_date or 'Atual'})"
            story.append(Paragraph(header, styles["BodyText"]))
            for highlight in experience.highlights:
                story.append(Paragraph(f"• {highlight}", styles["BodyText"]))
            story.append(Spacer(1, 6))

    if resume.educations:
        story.append(Paragraph("Formação", styles["Heading2"]))
        for education in resume.educations:
            text = f"<b>{education.degree}</b>"
            if education.institution:
                text += f" - {education.institution}"
            if education.summary:
                text += f" ({education.summary})"
            story.append(Paragraph(text, styles["BodyText"]))
        story.append(Spacer(1, 6))

    if resume.skills:
        story.append(Paragraph("Competências", styles["Heading2"]))
        story.append(Paragraph(", ".join(resume.skills), styles["BodyText"]))
        story.append(Spacer(1, 6))

    if resume.projects:
        story.append(Paragraph("Projetos", styles["Heading2"]))
        data = [["Projeto", "Descrição"]]
        for project in resume.projects:
            description = project.description or ""
            if project.link:
                description += f"\n{project.link}"
            data.append([project.name, description])
        table = Table(data, colWidths=[2.5 * inch, 3.5 * inch])
        table.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, 0), theme_color),
                    ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                    ("BOX", (0, 0), (-1, -1), 0.25, theme_color),
                    ("INNERGRID", (0, 0), (-1, -1), 0.25, theme_color),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ]
            )
        )
        story.append(table)

    doc.build(story)
    return buffer.getvalue()


def build_templates() -> Dict[str, TemplateDefinition]:
    moderno_azul = TemplateDefinition(
        metadata=TemplateMetadata(
            id="moderno-azul",
            name="Moderno Azul",
            description="Layout moderno com destaques em azul escuro",
            tags=["moderno", "profissional"],
        ),
        docx_renderer=_render_modern_blue_docx,
        pdf_renderer=lambda resume: _render_pdf_with_style(resume, colors.HexColor("#1F4E79"), "Moderno Azul"),
    )

    classico_serifado = TemplateDefinition(
        metadata=TemplateMetadata(
            id="classico-serifado",
            name="Clássico Serifado",
            description="Layout clássico com tipografia serifada",
            tags=["clássico", "formal"],
        ),
        docx_renderer=_render_classic_serif_docx,
        pdf_renderer=lambda resume: _render_pdf_with_style(resume, colors.HexColor("#333333"), "Clássico Serifado"),
    )

    minimalista_grade = TemplateDefinition(
        metadata=TemplateMetadata(
            id="minimalista-grade",
            name="Minimalista em Grade",
            description="Layout minimalista em duas colunas com blocos informativos",
            tags=["minimalista", "criativo"],
        ),
        docx_renderer=_render_minimalist_docx,
        pdf_renderer=lambda resume: _render_pdf_with_style(resume, colors.HexColor("#2E7D32"), "Minimalista em Grade"),
    )

    executivo_dourado = TemplateDefinition(
        metadata=TemplateMetadata(
            id="executivo-dourado",
            name="Executivo Dourado",
            description="Layout sofisticado com destaques em dourado e foco em resultados",
            tags=["executivo", "premium"],
        ),
        docx_renderer=_render_executive_gold_docx,
        pdf_renderer=lambda resume: _render_pdf_with_style(resume, colors.HexColor("#A57C00"), "Executivo Dourado"),
    )

    return {
        template.metadata.id: template
        for template in (moderno_azul, classico_serifado, minimalista_grade, executivo_dourado)
    }
